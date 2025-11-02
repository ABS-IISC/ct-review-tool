from flask import Flask, render_template_string, request, jsonify, send_file
import os
from werkzeug.utils import secure_filename
from datetime import datetime
import uuid
from docx import Document

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = '/tmp'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

sessions = {}

def extract_sections(doc):
    sections = {}
    current_section = "Main Content"
    content = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            is_header = False
            if para.runs and len(text) < 100:
                bold_runs = sum(1 for run in para.runs if run.bold)
                if bold_runs > len(para.runs) / 2:
                    is_header = True
            
            if is_header and (text.endswith(':') or text.isupper()):
                if content:
                    sections[current_section] = '\n'.join(content)
                current_section = text.rstrip(':')
                content = []
            else:
                content.append(text)
    
    if content:
        sections[current_section] = '\n'.join(content)
    
    return sections

def mock_analyze(section_name, content):
    return {
        "feedback_items": [{
            "id": "1",
            "type": "critical",
            "category": "Investigation Process",
            "description": f"Section '{section_name}' needs more detailed analysis according to Hawkeye guidelines.",
            "suggestion": "Add more specific details about the investigation methodology used.",
            "risk_level": "High",
            "confidence": 0.85
        }]
    }

HTML_TEMPLATE = '''
<!DOCTYPE html>
<html>
<head>
    <title>CT Review Tool</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 20px; background: #f5f5f5; }
        .container { max-width: 1200px; margin: 0 auto; background: white; padding: 20px; border-radius: 8px; }
        .header { text-align: center; margin-bottom: 30px; }
        .upload-area { border: 2px dashed #ccc; padding: 40px; text-align: center; margin-bottom: 20px; }
        .content { display: flex; gap: 20px; }
        .panel { flex: 1; border: 1px solid #ddd; border-radius: 4px; }
        .panel-header { background: #f8f9fa; padding: 10px; border-bottom: 1px solid #ddd; font-weight: bold; }
        .panel-content { padding: 15px; height: 400px; overflow-y: auto; }
        .feedback-item { background: #f8f9ff; border-left: 4px solid #007bff; padding: 15px; margin: 10px 0; }
        .btn { padding: 8px 16px; border: none; border-radius: 4px; cursor: pointer; margin: 5px; }
        .btn-primary { background: #007bff; color: white; }
        .btn-success { background: #28a745; color: white; }
        .btn-danger { background: #dc3545; color: white; }
        .hidden { display: none; }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>CT Review Tool</h1>
            <p>AI-Powered Document Analysis</p>
        </div>
        
        <div id="uploadSection">
            <div class="upload-area">
                <h3>Upload Document</h3>
                <input type="file" id="fileInput" accept=".docx" onchange="uploadFile()">
                <p>Select a .docx file to analyze</p>
            </div>
        </div>

        <div id="analysisSection" class="hidden">
            <div style="margin-bottom: 20px;">
                <select id="sectionSelect" onchange="loadSection()" style="padding: 8px; margin-right: 10px;">
                    <option value="">Select Section...</option>
                </select>
                <button class="btn btn-primary" onclick="completeReview()">Complete Review</button>
            </div>

            <div class="content">
                <div class="panel">
                    <div class="panel-header">Original Document</div>
                    <div class="panel-content" id="docContent">Select a section to view content...</div>
                </div>
                
                <div class="panel">
                    <div class="panel-header">AI Analysis</div>
                    <div class="panel-content" id="feedbackContent">Upload a document to begin analysis...</div>
                </div>
            </div>
        </div>
    </div>

    <script>
        let currentSession = null;
        let sections = [];

        function uploadFile() {
            const fileInput = document.getElementById('fileInput');
            const file = fileInput.files[0];
            if (!file) return;

            const formData = new FormData();
            formData.append('file', file);

            fetch('/upload', { method: 'POST', body: formData })
            .then(response => response.json())
            .then(data => {
                if (data.error) {
                    alert('Error: ' + data.error);
                } else {
                    currentSession = data.session_id;
                    sections = data.sections;
                    
                    const select = document.getElementById('sectionSelect');
                    select.innerHTML = '<option value="">Select Section...</option>';
                    sections.forEach((section, index) => {
                        const option = document.createElement('option');
                        option.value = index;
                        option.textContent = section;
                        select.appendChild(option);
                    });
                    
                    document.getElementById('uploadSection').classList.add('hidden');
                    document.getElementById('analysisSection').classList.remove('hidden');
                }
            });
        }

        function loadSection() {
            const select = document.getElementById('sectionSelect');
            const index = parseInt(select.value);
            if (isNaN(index)) return;

            document.getElementById('feedbackContent').innerHTML = 'Analyzing...';

            fetch(`/analyze/${currentSession}/${index}`)
            .then(response => response.json())
            .then(data => {
                if (data.error) {
                    document.getElementById('feedbackContent').innerHTML = 'Error: ' + data.error;
                } else {
                    document.getElementById('docContent').innerHTML = '<pre>' + data.content + '</pre>';
                    displayFeedback(data.feedback, data.section_name);
                }
            });
        }

        function displayFeedback(feedbackItems, sectionName) {
            const container = document.getElementById('feedbackContent');
            
            if (feedbackItems.length === 0) {
                container.innerHTML = '<p>No issues found in this section.</p>';
                return;
            }

            let html = '';
            feedbackItems.forEach((item, index) => {
                html += `
                    <div class="feedback-item">
                        <strong>${item.type.toUpperCase()}</strong> - ${item.risk_level} Risk<br>
                        <p>${item.description}</p>
                        ${item.suggestion ? '<p><em>' + item.suggestion + '</em></p>' : ''}
                        <button class="btn btn-success" onclick="acceptFeedback('${sectionName}', ${index})">Accept</button>
                        <button class="btn btn-danger" onclick="rejectFeedback(${index})">Reject</button>
                    </div>
                `;
            });
            
            container.innerHTML = html;
        }

        function acceptFeedback(sectionName, index) {
            fetch('/accept_feedback', {
                method: 'POST',
                headers: { 'Content-Type': 'application/json' },
                body: JSON.stringify({
                    session_id: currentSession,
                    section_name: sectionName,
                    feedback_item: { accepted: true, index: index }
                })
            });
            alert('Feedback accepted');
        }

        function rejectFeedback(index) {
            alert('Feedback rejected');
        }

        function completeReview() {
            if (!currentSession) return;
            
            fetch(`/complete_review/${currentSession}`)
            .then(response => response.blob())
            .then(blob => {
                const url = window.URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.href = url;
                a.download = 'reviewed_document.docx';
                a.click();
                window.URL.revokeObjectURL(url);
            });
        }
    </script>
</body>
</html>
'''

@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    if not file.filename.endswith('.docx'):
        return jsonify({'error': 'Please upload a .docx file'}), 400
    
    try:
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        doc = Document(filepath)
        sections = extract_sections(doc)
        
        session_id = str(uuid.uuid4())
        sessions[session_id] = {
            'filename': filename,
            'filepath': filepath,
            'sections': sections,
            'accepted_feedback': {}
        }
        
        return jsonify({
            'session_id': session_id,
            'sections': list(sections.keys()),
            'message': f'Loaded {len(sections)} sections'
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/analyze/<session_id>/<int:section_idx>')
def analyze_section(session_id, section_idx):
    if session_id not in sessions:
        return jsonify({'error': 'Session not found'}), 404
    
    session = sessions[session_id]
    section_names = list(session['sections'].keys())
    
    if section_idx >= len(section_names):
        return jsonify({'error': 'Section not found'}), 404
    
    section_name = section_names[section_idx]
    section_content = session['sections'][section_name]
    
    result = mock_analyze(section_name, section_content)
    
    return jsonify({
        'section_name': section_name,
        'content': section_content,
        'feedback': result.get('feedback_items', [])
    })

@app.route('/accept_feedback', methods=['POST'])
def accept_feedback():
    data = request.json
    session_id = data.get('session_id')
    section_name = data.get('section_name')
    feedback_item = data.get('feedback_item')
    
    if session_id not in sessions:
        return jsonify({'error': 'Session not found'}), 404
    
    session = sessions[session_id]
    
    if section_name not in session['accepted_feedback']:
        session['accepted_feedback'][section_name] = []
    
    session['accepted_feedback'][section_name].append(feedback_item)
    
    return jsonify({'message': 'Feedback accepted'})

@app.route('/complete_review/<session_id>')
def complete_review(session_id):
    if session_id not in sessions:
        return jsonify({'error': 'Session not found'}), 404
    
    session = sessions[session_id]
    
    try:
        doc = Document()
        doc.add_heading('CT Review Summary', 0)
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        
        for section_name, feedback_items in session['accepted_feedback'].items():
            doc.add_heading(f'Section: {section_name}', 1)
            for item in feedback_items:
                doc.add_paragraph(f"Feedback accepted for analysis point {item.get('index', 0) + 1}")
        
        output_path = f"/tmp/reviewed_{session['filename']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(output_path)
        
        return send_file(output_path, as_attachment=True)
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 8080)))