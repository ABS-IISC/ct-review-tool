"""
Core functionality for CT Review Tool - Business Logic Only
"""

import pandas as pd
import base64
import json
from datetime import datetime
import boto3
import threading
import os
import re
import traceback
import time
from pathlib import Path
import asyncio
import uuid
from collections import defaultdict
import queue
import zipfile
import shutil
from lxml import etree
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

# Global variables
guidelines_content = None
hawkeye_checklist = None
current_session = None
document_sections = {}
current_section_index = 0
accepted_feedback = defaultdict(list)
rejected_feedback = defaultdict(list)
user_feedback = defaultdict(list)
ai_feedback_cache = {}
current_section_feedback = []
review_completed = False
chat_history = []

# Define paths to guidelines documents
GUIDELINES_PATH = "CT_EE_Review_Guidelines.docx"
HAWKEYE_PATH = "Hawkeye_checklist.docx"

# Hawkeye checklist mapping
HAWKEYE_SECTIONS = {
    1: "Initial Assessment",
    2: "Investigation Process", 
    3: "Seller Classification",
    4: "Enforcement Decision-Making",
    5: "Additional Verification (High-Risk Cases)",
    6: "Multiple Appeals Handling",
    7: "Account Hijacking Prevention",
    8: "Funds Management",
    9: "REs-Q Outreach Process",
    10: "Sentiment Analysis",
    11: "Root Cause Analysis",
    12: "Preventative Actions",
    13: "Documentation and Reporting",
    14: "Cross-Team Collaboration",
    15: "Quality Control",
    16: "Continuous Improvement",
    17: "Communication Standards",
    18: "Performance Metrics",
    19: "Legal and Compliance",
    20: "New Service Launch Considerations"
}

# Standard writeup sections to look for
STANDARD_SECTIONS = [
    "Executive Summary",
    "Background",
    "Resolving Actions",
    "Root Cause",
    "Preventative Actions",
    "Investigation Process",
    "Seller Classification",
    "Documentation and Reporting",
    "Impact Assessment",
    "Timeline",
    "Recommendations"
]

# Sections to exclude from analysis
EXCLUDED_SECTIONS = [
    "Original Email",
    "Email Correspondence",
    "Raw Data",
    "Logs",
    "Attachments"
]

class WordDocumentWithComments:
    """Helper class to add comments to Word documents"""
    
    def __init__(self, doc_path):
        self.doc_path = doc_path
        self.temp_dir = f"temp_{uuid.uuid4()}"
        self.comments = []
        self.comment_id = 1
        
    def add_comment(self, paragraph_index, comment_text, author="AI Feedback"):
        """Add a comment to be inserted later"""
        self.comments.append({
            'id': self.comment_id,
            'paragraph_index': paragraph_index,
            'text': comment_text,
            'author': author,
            'date': datetime.now()
        })
        self.comment_id += 1
    
    def _create_comment_xml(self, comment):
        """Create comment XML structure"""
        comment_xml = f'''
        <w:comment w:id="{comment['id']}" w:author="{comment['author']}" 
                   w:date="{comment['date'].strftime('%Y-%m-%dT%H:%M:%S.%fZ')}" 
                   xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:p>
                <w:r>
                    <w:t>{comment['text']}</w:t>
                </w:r>
            </w:p>
        </w:comment>
        '''
        return comment_xml
    
    def save_with_comments(self, output_path):
        """Save document with comments added"""
        try:
            doc = Document(self.doc_path)
            temp_docx = f"{self.temp_dir}_temp.docx"
            doc.save(temp_docx)
            
            os.makedirs(self.temp_dir, exist_ok=True)
            with zipfile.ZipFile(temp_docx, 'r') as zip_ref:
                zip_ref.extractall(self.temp_dir)
            
            comments_path = os.path.join(self.temp_dir, 'word', 'comments.xml')
            
            comments_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
            <w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            '''
            
            for comment in self.comments:
                comments_xml += self._create_comment_xml(comment)
            
            comments_xml += '</w:comments>'
            
            with open(comments_path, 'w', encoding='utf-8') as f:
                f.write(comments_xml)
            
            # Update relationships and content types
            self._update_document_relationships()
            
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for root, dirs, files in os.walk(self.temp_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, self.temp_dir)
                        zipf.write(file_path, arcname)
            
            shutil.rmtree(self.temp_dir)
            os.remove(temp_docx)
            
            return True
            
        except Exception as e:
            print(f"Error adding comments: {str(e)}")
            if os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
            if os.path.exists(temp_docx):
                os.remove(temp_docx)
            return False

    def _update_document_relationships(self):
        """Update document relationships for comments"""
        rels_path = os.path.join(self.temp_dir, 'word', '_rels', 'document.xml.rels')
        if os.path.exists(rels_path):
            with open(rels_path, 'r', encoding='utf-8') as f:
                rels_content = f.read()
            
            if 'comments.xml' not in rels_content:
                new_rel = '<Relationship Id="rIdComments" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" Target="comments.xml"/>'
                rels_content = rels_content.replace('</Relationships>', f'{new_rel}</Relationships>')
                
                with open(rels_path, 'w', encoding='utf-8') as f:
                    f.write(rels_content)
        
        content_types_path = os.path.join(self.temp_dir, '[Content_Types].xml')
        if os.path.exists(content_types_path):
            with open(content_types_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if 'comments.xml' not in content:
                new_type = '<Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>'
                content = content.replace('</Types>', f'{new_type}</Types>')
                
                with open(content_types_path, 'w', encoding='utf-8') as f:
                    f.write(content)

class ReviewSession:
    def __init__(self):
        self.session_id = str(uuid.uuid4())
        self.start_time = datetime.now()
        self.document_name = ""
        self.document_content = ""
        self.document_object = None
        self.document_path = ""
        self.sections = {}
        self.section_paragraphs = {}
        self.paragraph_indices = {}
        self.current_section = 0
        self.feedback_history = defaultdict(list)
        self.section_status = {}

def load_guidelines():
    """Load the CT EE Review guidelines and Hawkeye checklist"""
    global guidelines_content, hawkeye_checklist
    
    try:
        if os.path.exists(GUIDELINES_PATH):
            guidelines_content = read_docx(GUIDELINES_PATH)
        
        if os.path.exists(HAWKEYE_PATH):
            hawkeye_checklist = read_docx(HAWKEYE_PATH)
        
        return guidelines_content, hawkeye_checklist
    except Exception as e:
        return None, None

def read_docx(file_path):
    """Extract text from a Word document"""
    try:
        doc = Document(file_path)
        full_text = []
        
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
                    
        return '\n'.join(full_text)
    except Exception as e:
        return f"Error reading document: {str(e)}"

def extract_document_sections_from_docx(doc):
    """Extract sections from Word document based on bold formatting"""
    sections = {}
    section_paragraphs = {}
    paragraph_indices = {}
    current_section = None
    current_content = []
    current_paragraphs = []
    current_indices = []
    
    for idx, para in enumerate(doc.paragraphs):
        is_bold = False
        if para.runs:
            bold_runs = sum(1 for run in para.runs if run.bold)
            total_runs = len(para.runs)
            is_bold = bold_runs > total_runs / 2
        
        text = para.text.strip()
        is_section_header = False
        
        if is_bold and text and len(text) < 100:
            for std_section in STANDARD_SECTIONS:
                if std_section.lower() in text.lower():
                    is_section_header = True
                    break
            
            if not is_section_header and (text.endswith(':') or text.isupper()):
                is_section_header = True
        
        if is_section_header:
            if current_section and current_content:
                exclude = False
                for excluded in EXCLUDED_SECTIONS:
                    if excluded.lower() in current_section.lower():
                        exclude = True
                        break
                
                if not exclude:
                    sections[current_section] = '\n'.join(current_content)
                    section_paragraphs[current_section] = current_paragraphs
                    paragraph_indices[current_section] = current_indices
            
            current_section = text.rstrip(':')
            current_content = []
            current_paragraphs = []
            current_indices = []
        else:
            if text:
                current_content.append(text)
                current_paragraphs.append(para)
                current_indices.append(idx)
    
    if current_section and current_content:
        exclude = False
        for excluded in EXCLUDED_SECTIONS:
            if excluded.lower() in current_section.lower():
                exclude = True
                break
        
        if not exclude:
            sections[current_section] = '\n'.join(current_content)
            section_paragraphs[current_section] = current_paragraphs
            paragraph_indices[current_section] = current_indices
    
    if not sections:
        all_text = []
        all_paras = []
        all_indices = []
        for idx, para in enumerate(doc.paragraphs):
            if para.text.strip():
                all_text.append(para.text)
                all_paras.append(para)
                all_indices.append(idx)
        sections = {"Main Content": '\n'.join(all_text)}
        section_paragraphs = {"Main Content": all_paras}
        paragraph_indices = {"Main Content": all_indices}
    
    return sections, section_paragraphs, paragraph_indices

def get_hawkeye_reference(category, content):
    """Map feedback to relevant Hawkeye checklist items"""
    references = []
    
    keyword_mapping = {
        1: ["customer experience", "cx impact", "customer trust", "buyer impact"],
        2: ["investigation", "sop", "enforcement decision", "abuse pattern"],
        3: ["seller classification", "good actor", "bad actor", "confused actor"],
        4: ["enforcement", "violation", "warning", "suspension"],
        5: ["verification", "supplier", "authenticity", "documentation"],
        6: ["appeal", "repeat", "retrospective"],
        7: ["hijacking", "security", "authentication", "secondary user"],
        8: ["funds", "disbursement", "financial"],
        9: ["outreach", "communication", "clarification"],
        10: ["sentiment", "escalation", "health safety", "legal threat"],
        11: ["root cause", "process gap", "system failure"],
        12: ["preventative", "solution", "improvement", "mitigation"],
        13: ["documentation", "reporting", "background"],
        14: ["cross-team", "collaboration", "engagement"],
        15: ["quality", "audit", "review", "performance"],
        16: ["continuous improvement", "training", "update"],
        17: ["communication standard", "messaging", "clarity"],
        18: ["metrics", "tracking", "measurement"],
        19: ["legal", "compliance", "regulation"],
        20: ["launch", "pilot", "rollback"]
    }
    
    content_lower = content.lower()
    category_lower = category.lower()
    
    for section_num, keywords in keyword_mapping.items():
        for keyword in keywords:
            if keyword in content_lower or keyword in category_lower:
                references.append({
                    'number': section_num,
                    'name': HAWKEYE_SECTIONS[section_num]
                })
                break
    
    return references[:3]

def classify_risk_level(feedback_item):
    """Classify risk level based on Hawkeye criteria"""
    high_risk_indicators = [
        "counterfeit", "fraud", "manipulation", "multiple violation",
        "immediate action", "legal", "health safety", "bad actor"
    ]
    
    medium_risk_indicators = [
        "pattern", "violation", "enforcement", "remediation",
        "correction", "warning"
    ]
    
    content_lower = f"{feedback_item.get('description', '')} {feedback_item.get('category', '')}".lower()
    
    for indicator in high_risk_indicators:
        if indicator in content_lower:
            return "High"
    
    for indicator in medium_risk_indicators:
        if indicator in content_lower:
            return "Medium"
    
    return "Low"

def invoke_aws_semantic_search(system_prompt, user_prompt, operation_name="LLM Analysis"):
    """AWS Bedrock invocation with Hawkeye guidelines"""
    global guidelines_content, hawkeye_checklist
    
    if guidelines_content is None or hawkeye_checklist is None:
        guidelines_content, hawkeye_checklist = load_guidelines()
    
    enhanced_system_prompt = system_prompt
    if hawkeye_checklist:
        truncated_hawkeye = hawkeye_checklist[:30000]
        enhanced_system_prompt = f"""{system_prompt}

HAWKEYE INVESTIGATION CHECKLIST:
{truncated_hawkeye}

Apply these Hawkeye investigation mental models in your analysis. Reference specific checklist items when providing feedback."""
    
    runtime = boto3.client('bedrock-runtime')
    
    body = json.dumps({
        "anthropic_version": "bedrock-2023-05-31",
        "max_tokens": 4000,
        "system": enhanced_system_prompt,
        "messages": [{"role": "user", "content": user_prompt}]
    })
    
    try:
        response = runtime.invoke_model(
            body=body,
            modelId='anthropic.claude-3-sonnet-20240229-v1:0',
            accept="application/json",
            contentType="application/json"
        )
        
        response_body = json.loads(response.get('body').read())
        return response_body['content'][0]['text']
        
    except Exception as e:
        # Return mock data for testing
        time.sleep(1)
        if "chat" in operation_name.lower():
            return "Based on the Hawkeye guidelines, I can help you understand the feedback better. The 20-point checklist emphasizes thorough investigation and customer impact assessment. What specific aspect would you like me to clarify?"
        
        return json.dumps({
            "feedback_items": [
                {
                    "id": "1",
                    "type": "critical",
                    "category": "investigation process",
                    "description": "Missing evaluation of customer experience (CX) impact. How might this abuse affect customer trust and satisfaction?",
                    "suggestion": "Add analysis of potential negative reviews, returns, or complaints that could result from this issue",
                    "example": "Consider both immediate and long-term effects on customer trust as outlined in Hawkeye #1",
                    "questions": [
                        "Have you evaluated the customer experience (CX) impact?",
                        "Did you consider how this affects buyer trust?"
                    ],
                    "confidence": 0.95
                }
            ]
        })

def analyze_section_with_ai(section_name, section_content, doc_type="Full Write-up"):
    """Analyze a single section with Hawkeye framework"""
    
    cache_key = f"{section_name}_{hash(section_content)}"
    if cache_key in ai_feedback_cache:
        return ai_feedback_cache[cache_key]
    
    prompt = f"""Analyze this section "{section_name}" from a {doc_type} document using the Hawkeye investigation framework.

SECTION CONTENT:
{section_content[:3000]}

Provide feedback following the 20-point Hawkeye checklist. For each feedback item, include:
1. Specific questions from the Hawkeye checklist that should be addressed
2. References to relevant Hawkeye checkpoint numbers (#1-20)
3. Examples from the case studies when applicable
4. Risk classification (High/Medium/Low)

Return feedback in this JSON format:
{{
    "feedback_items": [
        {{
            "id": "unique_id",
            "type": "critical|important|suggestion|positive",
            "category": "category matching Hawkeye sections",
            "description": "Clear description referencing Hawkeye criteria",
            "suggestion": "Specific suggestion based on Hawkeye guidelines",
            "example": "Example from case studies or Hawkeye checklist",
            "questions": ["Question 1 from Hawkeye?", "Question 2?"],
            "hawkeye_refs": [1, 11, 12],
            "risk_level": "High|Medium|Low",
            "confidence": 0.95
        }}
    ]
}}"""
    
    system_prompt = "You are an expert document reviewer following the Hawkeye investigation mental models for CT EE guidelines."
    
    response = invoke_aws_semantic_search(system_prompt, prompt, f"Hawkeye Analysis: {section_name}")
    
    try:
        result = json.loads(response)
    except:
        json_match = re.search(r'\{.*\}', response, re.DOTALL)
        if json_match:
            try:
                result = json.loads(json_match.group(0))
            except:
                result = {"feedback_items": []}
        else:
            result = {"feedback_items": []}
    
    for item in result.get('feedback_items', []):
        if 'hawkeye_refs' not in item:
            refs = get_hawkeye_reference(item.get('category', ''), item.get('description', ''))
            item['hawkeye_refs'] = [ref['number'] for ref in refs]
        
        if 'risk_level' not in item:
            item['risk_level'] = classify_risk_level(item)
    
    ai_feedback_cache[cache_key] = result
    return result

def process_chat_query(query, context):
    """Process chat query with context awareness"""
    global current_session, hawkeye_checklist, current_section_feedback
    
    context_info = f"""
    Current Section: {context.get('current_section', 'None')}
    Current Feedback Items: {len(current_section_feedback)}
    Document Type: Full Write-up
    """
    
    if current_section_feedback:
        context_info += "\nCurrent Section Feedback Summary:\n"
        for item in current_section_feedback[:3]:
            context_info += f"- {item['type']}: {item['description'][:100]}...\n"
    
    prompt = f"""You are an AI assistant helping with document review using the Hawkeye framework.

CONTEXT:
{context_info}

HAWKEYE GUIDELINES REFERENCE:
The 20-point Hawkeye checklist includes:
1. Initial Assessment - Evaluate CX impact
2. Investigation Process - Challenge SOPs
3. Seller Classification - Identify good/bad actors
4. Enforcement Decision-Making
5. Additional Verification for High-Risk Cases
...and 15 more points

USER QUESTION: {query}

Provide a helpful, specific response that references the Hawkeye guidelines when relevant. Be concise but thorough."""
    
    system_prompt = "You are an expert assistant for the Hawkeye document review system."
    
    response = invoke_aws_semantic_search(system_prompt, prompt, "Chat Assistant")
    
    return response

def create_reviewed_document_with_proper_comments(original_doc_path, doc_name, comments_data):
    """Create a copy of the original document with proper Word comments"""
    
    try:
        doc_with_comments = WordDocumentWithComments(original_doc_path)
        
        for comment_data in comments_data:
            author = comment_data.get('author', 'AI Feedback')
            doc_with_comments.add_comment(
                paragraph_index=comment_data['paragraph_index'],
                comment_text=comment_data['comment'],
                author=author
            )
        
        output_path = f'reviewed_{doc_name}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        success = doc_with_comments.save_with_comments(output_path)
        
        if success:
            return output_path
        else:
            return create_simple_reviewed_copy(original_doc_path, doc_name, comments_data)
            
    except Exception as e:
        print(f"Error creating document with comments: {str(e)}")
        return create_simple_reviewed_copy(original_doc_path, doc_name, comments_data)

def create_simple_reviewed_copy(original_doc_path, doc_name, comments_data):
    """Create a simple copy with inline comment markers as fallback"""
    try:
        output_path = f'reviewed_{doc_name}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        
        doc = Document(original_doc_path)
        
        doc.add_page_break()
        heading = doc.add_heading('Hawkeye Review Feedback Summary', 1)
        
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        doc.add_paragraph(f'Total feedback items: {len(comments_data)}')
        doc.add_paragraph('')
        
        section_comments = defaultdict(list)
        for comment in comments_data:
            section_comments[comment['section']].append(comment)
        
        for section, comments in section_comments.items():
            section_heading = doc.add_heading(section, 2)
            
            for comment in comments:
                p = doc.add_paragraph(style='List Bullet')
                author = comment.get('author', 'AI Feedback')
                p.add_run(f"[{author}] {comment['type'].upper()} - {comment['risk_level']} Risk: ").bold = True
                p.add_run(comment['comment'])
        
        doc.save(output_path)
        return output_path
        
    except Exception as e:
        print(f"Error creating simple copy: {str(e)}")
        return None